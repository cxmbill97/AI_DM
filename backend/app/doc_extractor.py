"""Document text extraction for script ingestion.

Supports PDF, DOCX, and plain-text TXT files.
Pure function — no LLM calls, no async IO.
"""

from __future__ import annotations

import io
from pathlib import Path


class UnsupportedFormatError(ValueError):
    """Raised when the uploaded file extension is not supported."""


class ExtractionError(RuntimeError):
    """Raised when the underlying library fails to read the file."""


def extract_text(filename: str, content: bytes) -> str:
    """Extract plain text from PDF, DOCX, or TXT bytes.

    Parameters
    ----------
    filename:
        Original filename — extension determines the parser.
    content:
        Raw file bytes.

    Returns
    -------
    str
        Extracted text with sections joined by double newlines.

    Raises
    ------
    UnsupportedFormatError
        If the extension is not .pdf, .docx, or .txt.
    ExtractionError
        If the underlying library fails (corrupt file, password-protected PDF, etc.)
    """
    ext = Path(filename).suffix.lower()
    if ext == ".txt":
        return _extract_txt(content)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    raise UnsupportedFormatError(f"Unsupported file format: {ext!r}. Please upload PDF, DOCX, or TXT.")


# ---------------------------------------------------------------------------
# Format-specific helpers
# ---------------------------------------------------------------------------


def _extract_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="replace").strip()


def _extract_pdf(content: bytes) -> str:
    try:
        import pdfplumber  # noqa: PLC0415
    except ImportError as e:
        raise ExtractionError("pdfplumber is not installed") from e

    try:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(layout=True) or ""
                if text.strip():
                    pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as e:
        raise ExtractionError(f"PDF extraction failed: {e}") from e


def _extract_docx(content: bytes) -> str:
    try:
        import docx  # noqa: PLC0415
    except ImportError as e:
        raise ExtractionError("python-docx is not installed") from e

    try:
        doc = docx.Document(io.BytesIO(content))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables (tab-separated cells, rows joined by newline)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        return "\n".join(parts)
    except Exception as e:
        raise ExtractionError(f"DOCX extraction failed: {e}") from e
